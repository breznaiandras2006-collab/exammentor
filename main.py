# main.py
from __future__ import annotations

import hashlib
import io
import os
import random
from pathlib import Path
from urllib.parse import quote_plus

from fastapi import FastAPI, Request, UploadFile, File, Form
from fastapi.responses import (
    HTMLResponse,
    RedirectResponse,
    FileResponse,
    PlainTextResponse,
    StreamingResponse,
)
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from db import (
    init_db,
    list_documents,
    get_document,
    insert_document,
    search_documents,
    list_notes,
    get_note,
    insert_note,
    update_note,
    delete_note,
    search_notes,
    get_all_settings,
    set_setting,
    # Study
    create_study_card,
    update_study_card,
    delete_study_card,
    get_study_card,
    list_study_cards,
    get_study_counts,
    get_next_due_card,
    get_random_card,
    get_random_distractors,
    review_card,
    study_stats,
)
from tools import (
    safe_filename,
    make_snippet,
    extract_text_from_pdf,
    pdf_page_count,
    compress_pdf_bytes,
    split_pdf_bytes,
    parse_ranges,
    merge_pdf_bytes,
    delete_pages_pdf_bytes,
    rotate_pages_pdf_bytes,
    parse_page_sequence,
    extract_pages_pdf_bytes,
    reorder_pages_pdf_bytes,
    # Study
    extract_qa_pairs,
)

app = FastAPI()

BASE_DIR = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))

# static (css, js, etc.)
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")


@app.on_event("startup")
def _startup():
    init_db()
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def _settings_context() -> dict:
    """Templates can use: s.answer_language, s.theme, etc."""
    s = get_all_settings() or {}
    s2 = dict(s)
    # IMPORTANT: strings like "0" are truthy in Jinja â†’ convert the common bools
    s2["manual_mode"] = (s.get("manual_mode", "0") == "1")
    return {"s": s2}


def _parse_doc_ids(text_value: str):
    """Parse comma-separated doc ids: "1,2,3" -> [1,2,3]."""
    raw = (text_value or "").strip()
    if not raw:
        return []
    out = []
    for part in raw.split(","):
        p = part.strip()
        if not p:
            continue
        try:
            out.append(int(p))
        except Exception:
            continue
    # unique, keep order
    seen = set()
    uniq = []
    for i in out:
        if i in seen:
            continue
        seen.add(i)
        uniq.append(i)
    return uniq


def _store_pdf_bytes_as_document(
    pdf_bytes: bytes,
    *,
    title: str,
    original_name: str,
    language: str = "auto",
) -> int:
    """Save generated PDF into uploads/ + insert into documents table."""
    safe_orig = safe_filename(original_name, "document.pdf")
    raw = (safe_orig + str(os.urandom(8))).encode("utf-8", "ignore")
    h = hashlib.sha256(raw).hexdigest()[:24]
    stored_name = f"{h}_{safe_orig}"

    (UPLOAD_DIR / stored_name).write_bytes(pdf_bytes)

    try:
        pages = pdf_page_count(pdf_bytes)
    except Exception:
        pages = 0

    try:
        search_text = extract_text_from_pdf(pdf_bytes, max_pages=25)
    except Exception:
        search_text = ""

    return insert_document(
        title=title,
        original_name=original_name,
        stored_name=stored_name,
        language=(language or "auto"),
        pages=pages,
        doc_type="pdf",
        search_text=search_text,
    )


# ---------------- Home / Onboarding ----------------
@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    ctx = {"request": request}
    ctx.update(_settings_context())
    return templates.TemplateResponse("index.html", ctx)


@app.get("/onboarding", response_class=HTMLResponse)
def onboarding(request: Request):
    ctx = {"request": request}
    ctx.update(_settings_context())
    return templates.TemplateResponse("onboarding.html", ctx)


# ---------------- Documents ----------------
@app.get("/documents", response_class=HTMLResponse)
def documents_page(request: Request):
    docs = list_documents()
    ctx = {"request": request, "docs": docs}
    ctx.update(_settings_context())
    return templates.TemplateResponse("documents.html", ctx)


@app.post("/documents/upload")
async def documents_upload(
    request: Request,
    title: str = Form(""),
    language: str = Form("auto"),
    pdf: UploadFile = File(...),
):
    original = pdf.filename or "document.pdf"
    safe_orig = safe_filename(original, "document.pdf")

    raw = (original + str(os.urandom(8))).encode("utf-8", "ignore")
    h = hashlib.sha256(raw).hexdigest()[:24]
    stored_name = f"{h}_{safe_orig}"

    target = UPLOAD_DIR / stored_name
    data = await pdf.read()
    target.write_bytes(data)

    # quick extraction for text PDFs
    try:
        pages = pdf_page_count(data)
    except Exception:
        pages = 0

    try:
        search_text = extract_text_from_pdf(data, max_pages=25)
    except Exception:
        search_text = ""

    title2 = (title or "").strip() or Path(original).stem

    doc_id = insert_document(
        title=title2,
        original_name=original,
        stored_name=stored_name,
        language=(language or "auto"),
        pages=pages,
        doc_type="pdf",
        search_text=search_text,
    )

    return RedirectResponse(url=f"/documents/{doc_id}", status_code=303)


@app.get("/documents/{doc_id}", response_class=HTMLResponse)
def document_detail(request: Request, doc_id: int):
    doc = get_document(doc_id)
    if not doc:
        ctx = {"request": request, "message": "Document not found"}
        ctx.update(_settings_context())
        return templates.TemplateResponse("not_found.html", ctx, status_code=404)

    file_url = f"/documents/{doc_id}/file"
    ctx = {"request": request, "doc": doc, "file_url": file_url}
    ctx.update(_settings_context())
    return templates.TemplateResponse("document_detail.html", ctx)


@app.get("/documents/{doc_id}/file")
def document_file(doc_id: int):
    doc = get_document(doc_id)
    if not doc:
        return PlainTextResponse("Not found", status_code=404)

    fp = UPLOAD_DIR / (doc.get("stored_name") or "")
    if not fp.exists():
        return PlainTextResponse("File missing", status_code=404)

    # inline -> open in browser (not forced download)
    return FileResponse(
        str(fp),
        media_type="application/pdf",
        filename=doc.get("original_name", "document.pdf"),
        headers={
            "Content-Disposition": f'inline; filename="{doc.get("original_name","document.pdf")}"'
        },
    )


@app.get("/documents/{doc_id}/download")
def document_download(doc_id: int):
    doc = get_document(doc_id)
    if not doc:
        return PlainTextResponse("Not found", status_code=404)

    fp = UPLOAD_DIR / (doc.get("stored_name") or "")
    if not fp.exists():
        return PlainTextResponse("File missing", status_code=404)

    return FileResponse(
        str(fp),
        media_type="application/pdf",
        filename=doc.get("original_name", "document.pdf"),
        headers={
            "Content-Disposition": f'attachment; filename="{doc.get("original_name","document.pdf")}"'
        },
    )


# ---------------- Notes ----------------
@app.get("/notes", response_class=HTMLResponse)
def notes_page(request: Request, doc: str = ""):
    docs = list_documents()

    doc_selected = None
    if (doc or "").strip():
        try:
            doc_selected = int(doc)
        except Exception:
            doc_selected = None

    notes = list_notes(document_id=doc_selected)

    ctx = {
        "request": request,
        "notes": notes,
        "docs": docs,
        "doc_selected": doc_selected,
    }
    ctx.update(_settings_context())
    return templates.TemplateResponse("notes.html", ctx)


@app.get("/notes/new", response_class=HTMLResponse)
def note_new_page(request: Request, doc: str = ""):
    doc_selected = None
    if (doc or "").strip():
        try:
            doc_selected = int(doc)
        except Exception:
            doc_selected = None

    ctx = {
        "request": request,
        "docs": list_documents(),
        "mode": "new",
        "note": {"title": "", "body": "", "document_id": doc_selected},
    }
    ctx.update(_settings_context())
    return templates.TemplateResponse("note_edit.html", ctx)


@app.post("/notes/new")
def note_new_post(
    request: Request,
    title: str = Form(""),
    body: str = Form(""),
    document_id: str = Form(""),
):
    title2 = (title or "").strip() or "Untitled"
    body2 = (body or "").strip()

    doc_id_val = None
    if (document_id or "").strip():
        try:
            doc_id_val = int(document_id)
        except Exception:
            doc_id_val = None

    nid = insert_note(title2, body2, doc_id_val)
    return RedirectResponse(url=f"/notes/{nid}", status_code=303)


@app.get("/notes/{note_id}", response_class=HTMLResponse)
def note_detail(request: Request, note_id: int):
    note = get_note(note_id)
    if not note:
        ctx = {"request": request, "message": "Note not found"}
        ctx.update(_settings_context())
        return templates.TemplateResponse("not_found.html", ctx, status_code=404)

    doc = None
    if note.get("document_id"):
        doc = get_document(int(note["document_id"]))

    ctx = {"request": request, "note": note, "doc": doc}
    ctx.update(_settings_context())
    return templates.TemplateResponse("note_detail.html", ctx)


@app.get("/notes/{note_id}/edit", response_class=HTMLResponse)
def note_edit_page(request: Request, note_id: int):
    note = get_note(note_id)
    if not note:
        ctx = {"request": request, "message": "Note not found"}
        ctx.update(_settings_context())
        return templates.TemplateResponse("not_found.html", ctx, status_code=404)

    ctx = {
        "request": request,
        "note": note,
        "docs": list_documents(),
        "mode": "edit",
    }
    ctx.update(_settings_context())
    return templates.TemplateResponse("note_edit.html", ctx)


@app.post("/notes/{note_id}/edit")
def note_edit_post(
    request: Request,
    note_id: int,
    title: str = Form(""),
    body: str = Form(""),
    document_id: str = Form(""),
):
    doc_id_val = None
    if (document_id or "").strip():
        try:
            doc_id_val = int(document_id)
        except Exception:
            doc_id_val = None

    update_note(note_id, (title or "").strip() or "Untitled", (body or "").strip(), doc_id_val)
    return RedirectResponse(url=f"/notes/{note_id}", status_code=303)


@app.post("/notes/{note_id}/delete")
def note_delete(request: Request, note_id: int):
    delete_note(note_id)
    return RedirectResponse(url="/notes", status_code=303)


@app.get("/notes/{note_id}/export/docx")
def note_export_docx(note_id: int):
    note = get_note(note_id)
    if not note:
        return PlainTextResponse("Not found", status_code=404)

    docx = DocxDocument()
    docx.add_heading(note.get("title") or "Note", level=1)

    body = (note.get("body") or "").splitlines()
    for line in body:
        # keep it simple (later: markdown -> rich)
        docx.add_paragraph(line)

    buf = io.BytesIO()
    docx.save(buf)
    buf.seek(0)

    filename = safe_filename((note.get("title") or "note") + ".docx", "note.docx")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/notes/{note_id}/export/pdf")
def note_export_pdf(note_id: int):
    note = get_note(note_id)
    if not note:
        return PlainTextResponse("Not found", status_code=404)

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    title = note.get("title") or "Note"
    c.setFont("Helvetica-Bold", 16)
    c.drawString(48, height - 64, title)

    c.setFont("Helvetica", 11)
    y = height - 92
    for line in (note.get("body") or "").splitlines():
        # basic line-wrapping
        text = line.rstrip()
        if not text:
            y -= 14
            continue
        while len(text) > 110:
            c.drawString(48, y, text[:110])
            text = text[110:]
            y -= 14
            if y < 60:
                c.showPage()
                c.setFont("Helvetica", 11)
                y = height - 64
        c.drawString(48, y, text)
        y -= 14
        if y < 60:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - 64

    c.showPage()
    c.save()
    buf.seek(0)

    filename = safe_filename((note.get("title") or "note") + ".pdf", "note.pdf")
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------- Ask ----------------
@app.get("/ask", response_class=HTMLResponse)
def ask_get(request: Request, q: str = "", scope: str = "all", doc: str = ""):
    docs_list = list_documents()

    doc_id_val = None
    if (doc or "").strip():
        try:
            doc_id_val = int(doc)
        except Exception:
            doc_id_val = None

    q2 = (q or "").strip()
    results_notes = []
    results_docs = []
    answer_lines = []

    if q2:
        if scope in ("all", "notes"):
            results_notes = search_notes(q=q2, document_id=doc_id_val, limit=12)
        if scope in ("all", "docs"):
            results_docs = search_documents(q=q2, limit=8)

        # precompute snippets for templates
        for n in results_notes:
            n['snippet'] = make_snippet(n.get('body',''), q2)
        for d in results_docs:
            d['snippet'] = make_snippet(d.get('search_text',''), q2)

        for n in results_notes[:6]:
            snippet = make_snippet(n.get("body", ""), q2)
            answer_lines.append(f"ðŸ“ {n.get('title','')} â€” {snippet}")

        for d in results_docs[:4]:
            snippet = make_snippet(d.get("search_text", ""), q2)
            answer_lines.append(f"ðŸ“„ {d.get('title','')} â€” {snippet}")

    compiled_answer = "\n".join(answer_lines).strip()

    ctx = {
        "request": request,
        "q": q2,
        "scope": scope,
        "docs": docs_list,
        "doc_selected": doc_id_val,
        "results_notes": results_notes,
        "results_docs": results_docs,
        "compiled_answer": compiled_answer,
    }
    ctx.update(_settings_context())
    return templates.TemplateResponse("ask.html", ctx)


@app.post("/ask")
def ask_post(request: Request, q: str = Form(""), scope: str = Form("all"), doc: str = Form("")):
    q_enc = quote_plus((q or "").strip())
    scope_enc = quote_plus((scope or "all").strip())
    doc_enc = quote_plus((doc or "").strip())
    return RedirectResponse(url=f"/ask?q={q_enc}&scope={scope_enc}&doc={doc_enc}", status_code=303)


@app.post("/ask/save-note")
def ask_save_note(
    request: Request,
    title: str = Form(""),
    body: str = Form(""),
    document_id: str = Form(""),
):
    title2 = (title or "").strip() or "Ask result"
    body2 = (body or "").strip()

    doc_id_val = None
    if (document_id or "").strip():
        try:
            doc_id_val = int(document_id)
        except Exception:
            doc_id_val = None

    nid = insert_note(title2, body2, doc_id_val)
    return RedirectResponse(url=f"/notes/{nid}", status_code=303)


# ---------------- Settings ----------------
@app.get("/settings", response_class=HTMLResponse)
def settings_page(request: Request):
    ctx = {"request": request}
    ctx.update(_settings_context())
    return templates.TemplateResponse("settings.html", ctx)


@app.post("/settings")
def settings_save(
    request: Request,
    ui_lang: str = Form("hu"),
    answer_language: str = Form("hu"),
    theme: str = Form("dark"),
    manual_mode: str = Form(""),
    translation_style: str = Form("precise"),
    default_gpt_mode: str = Form("exam"),
):
    set_setting("ui_lang", (ui_lang or "hu").strip())
    set_setting("answer_language", (answer_language or "hu").strip())
    set_setting("theme", (theme or "dark").strip())

    # checkbox returns "1" or missing
    set_setting("manual_mode", "1" if (manual_mode == "1") else "0")
    set_setting("translation_style", (translation_style or "precise").strip())
    set_setting("default_gpt_mode", (default_gpt_mode or "exam").strip())

    return RedirectResponse(url="/settings", status_code=303)


# ---------------- PDF Tools (B modul) ----------------
def _render_pdf_tools(
    request: Request,
    *,
    status_code: int = 200,
    result: str | None = None,
    result_kind: str = "info",  # info|success|warn|error
    created_docs: list | None = None,
    form: dict | None = None,
):
    """Consistent PDF Tools page rendering with basic UX feedback."""
    ctx = {
        "request": request,
        "docs": list_documents(),
        "result": result,
        "result_kind": result_kind,
        "created_docs": created_docs or [],
        "form": form or {},
    }
    ctx.update(_settings_context())
    return templates.TemplateResponse("pdf_tools.html", ctx, status_code=status_code)


@app.get("/pdf-tools", response_class=HTMLResponse)
def pdf_tools_page(request: Request):
    return _render_pdf_tools(request)


@app.post("/pdf-tools/compress", response_class=HTMLResponse)
def pdf_tools_compress(request: Request, doc_id: int = Form(...)):
    doc = get_document(doc_id)
    if not doc:
        return _render_pdf_tools(
            request,
            status_code=404,
            result="âŒ Nem talÃ¡lom ezt a dokumentumot.",
            result_kind="error",
            form={"compress_doc_id": doc_id},
        )

    fp = UPLOAD_DIR / (doc.get("stored_name") or "")
    if not fp.exists():
        return _render_pdf_tools(
            request,
            status_code=404,
            result="âŒ A PDF fÃ¡jl nem talÃ¡lhatÃ³ a szerveren (hiÃ¡nyzÃ³ uploads fÃ¡jl).",
            result_kind="error",
            form={"compress_doc_id": doc_id},
        )

    data = fp.read_bytes()
    out = compress_pdf_bytes(data)

    original_out = f"compressed_{doc.get('original_name','document.pdf')}"
    title_out = f"Compressed â€” {doc.get('title') or doc.get('original_name','PDF')}"
    new_id = _store_pdf_bytes_as_document(out, title=title_out, original_name=original_out, language=doc.get("language") or "auto")

    new_doc = get_document(new_id)
    return _render_pdf_tools(
        request,
        result="âœ… TÃ¶mÃ¶rÃ­tett PDF elkÃ©szÃ¼lt.",
        result_kind="success",
        created_docs=[new_doc] if new_doc else [],
        form={"compress_doc_id": doc_id},
    )


@app.post("/pdf-tools/split", response_class=HTMLResponse)
def pdf_tools_split(request: Request, doc_id: int = Form(...), ranges_text: str = Form("")):
    doc = get_document(doc_id)
    if not doc:
        return _render_pdf_tools(
            request,
            status_code=404,
            result="âŒ Nem talÃ¡lom ezt a dokumentumot.",
            result_kind="error",
            form={"split_doc_id": doc_id, "split_ranges_text": ranges_text},
        )

    fp = UPLOAD_DIR / (doc.get("stored_name") or "")
    if not fp.exists():
        return _render_pdf_tools(
            request,
            status_code=404,
            result="âŒ A PDF fÃ¡jl nem talÃ¡lhatÃ³ a szerveren (hiÃ¡nyzÃ³ uploads fÃ¡jl).",
            result_kind="error",
            form={"split_doc_id": doc_id, "split_ranges_text": ranges_text},
        )

    try:
        ranges = parse_ranges(ranges_text)
    except Exception as e:
        return _render_pdf_tools(
            request,
            status_code=400,
            result=f"âš ï¸ HibÃ¡s tartomÃ¡ny. PÃ©lda: 1-3, 5, 8-10. RÃ©szlet: {e}",
            result_kind="warn",
            form={"split_doc_id": doc_id, "split_ranges_text": ranges_text},
        )
    if not ranges:
        return _render_pdf_tools(
            request,
            status_code=400,
            result="âš ï¸ Adj meg oldaltartomÃ¡nyt (pÃ©lda: 1-2, 4 vagy 1-3;5).",
            result_kind="warn",
            form={"split_doc_id": doc_id, "split_ranges_text": ranges_text},
        )

    data = fp.read_bytes()
    parts = split_pdf_bytes(data, ranges)

    created = []
    base = Path(doc.get("original_name") or "document.pdf").stem
    for fname, bts in parts:
        out_original = f"{base}_{fname}"
        out_title = f"{doc.get('title') or base} â€” {fname.replace('.pdf','')}"
        new_id = _store_pdf_bytes_as_document(bts, title=out_title, original_name=out_original, language=doc.get("language") or "auto")
        nd = get_document(new_id)
        if nd:
            created.append(nd)

    msg = f"âœ… Split kÃ©sz: {len(created)} rÃ©sz."
    kind = "success" if created else "warn"
    if not created:
        msg = "âš ï¸ Nem jÃ¶tt lÃ©tre rÃ©sz PDF (ellenÅ‘rizd az oldaltartomÃ¡nyt)."
    return _render_pdf_tools(
        request,
        result=msg,
        result_kind=kind,
        created_docs=created,
        form={"split_doc_id": doc_id, "split_ranges_text": ranges_text},
    )


@app.post("/pdf-tools/merge", response_class=HTMLResponse)
def pdf_tools_merge(request: Request, doc_ids_text: str = Form("")):
    ids = _parse_doc_ids(doc_ids_text)
    if len(ids) < 2:
        return _render_pdf_tools(
            request,
            status_code=400,
            result="âš ï¸ Adj meg legalÃ¡bb 2 doc ID-t (pÃ©lda: 1,2).",
            result_kind="warn",
            form={"merge_doc_ids_text": doc_ids_text},
        )

    pdf_list = []
    titles = []
    lang = "auto"

    for doc_id in ids:
        doc = get_document(doc_id)
        if not doc:
            continue
        fp = UPLOAD_DIR / (doc.get("stored_name") or "")
        if not fp.exists():
            continue
        pdf_list.append(fp.read_bytes())
        titles.append(doc.get("title") or f"#{doc_id}")
        lang = doc.get("language") or lang

    if len(pdf_list) < 2:
        return _render_pdf_tools(
            request,
            status_code=404,
            result="âŒ Nem talÃ¡ltam legalÃ¡bb 2 Ã©rvÃ©nyes PDF-et a megadott ID-khez.",
            result_kind="error",
            form={"merge_doc_ids_text": doc_ids_text},
        )

    out = merge_pdf_bytes(pdf_list)
    out_original = f"merged_{'_'.join(str(i) for i in ids)}.pdf"
    out_title = "Merged â€” " + ", ".join(titles[:3]) + ("â€¦" if len(titles) > 3 else "")

    new_id = _store_pdf_bytes_as_document(out, title=out_title, original_name=out_original, language=lang)
    new_doc = get_document(new_id)

    return _render_pdf_tools(
        request,
        result="âœ… Ã–sszefÅ±zÃ¶tt PDF elkÃ©szÃ¼lt.",
        result_kind="success",
        created_docs=[new_doc] if new_doc else [],
        form={"merge_doc_ids_text": doc_ids_text},
    )


@app.post("/pdf-tools/delete-pages", response_class=HTMLResponse)
def pdf_tools_delete_pages(request: Request, doc_id: int = Form(...), ranges_text: str = Form("")):
    doc = get_document(doc_id)
    if not doc:
        return _render_pdf_tools(
            request,
            status_code=404,
            result="âŒ Nem talÃ¡lom ezt a dokumentumot.",
            result_kind="error",
            form={"delete_doc_id": doc_id, "delete_ranges_text": ranges_text},
        )

    try:
        ranges = parse_ranges(ranges_text)
    except Exception as e:
        return _render_pdf_tools(
            request,
            status_code=400,
            result=f"âš ï¸ HibÃ¡s tartomÃ¡ny. PÃ©lda: 1-2, 4. RÃ©szlet: {e}",
            result_kind="warn",
            form={"delete_doc_id": doc_id, "delete_ranges_text": ranges_text},
        )
    if not ranges:
        return _render_pdf_tools(
            request,
            status_code=400,
            result="âš ï¸ Adj meg oldaltartomÃ¡nyt (pÃ©lda: 1-2, 4).",
            result_kind="warn",
            form={"delete_doc_id": doc_id, "delete_ranges_text": ranges_text},
        )

    fp = UPLOAD_DIR / (doc.get("stored_name") or "")
    if not fp.exists():
        return _render_pdf_tools(
            request,
            status_code=404,
            result="âŒ A PDF fÃ¡jl nem talÃ¡lhatÃ³ a szerveren (hiÃ¡nyzÃ³ uploads fÃ¡jl).",
            result_kind="error",
            form={"delete_doc_id": doc_id, "delete_ranges_text": ranges_text},
        )

    out = delete_pages_pdf_bytes(fp.read_bytes(), ranges)
    out_original = f"pages_removed_{doc.get('original_name','document.pdf')}"
    out_title = f"Pages removed â€” {doc.get('title') or doc.get('original_name','PDF')}"

    new_id = _store_pdf_bytes_as_document(out, title=out_title, original_name=out_original, language=doc.get("language") or "auto")
    new_doc = get_document(new_id)

    return _render_pdf_tools(
        request,
        result="âœ… Oldalak tÃ¶rÃ¶lve (Ãºj dokumentum kÃ©szÃ¼lt).",
        result_kind="success",
        created_docs=[new_doc] if new_doc else [],
        form={"delete_doc_id": doc_id, "delete_ranges_text": ranges_text},
    )


@app.post("/pdf-tools/rotate", response_class=HTMLResponse)
def pdf_tools_rotate(request: Request, doc_id: int = Form(...), ranges_text: str = Form(""), degrees: int = Form(90)):
    doc = get_document(doc_id)
    if not doc:
        return _render_pdf_tools(
            request,
            status_code=404,
            result="âŒ Nem talÃ¡lom ezt a dokumentumot.",
            result_kind="error",
            form={"rotate_doc_id": doc_id, "rotate_ranges_text": ranges_text, "rotate_degrees": degrees},
        )

    try:
        ranges = parse_ranges(ranges_text)
    except Exception as e:
        return _render_pdf_tools(
            request,
            status_code=400,
            result=f"âš ï¸ HibÃ¡s tartomÃ¡ny. PÃ©lda: 1-2, 4. RÃ©szlet: {e}",
            result_kind="warn",
            form={"rotate_doc_id": doc_id, "rotate_ranges_text": ranges_text, "rotate_degrees": degrees},
        )
    if not ranges:
        return _render_pdf_tools(
            request,
            status_code=400,
            result="âš ï¸ Adj meg oldaltartomÃ¡nyt (pÃ©lda: 1-2, 4).",
            result_kind="warn",
            form={"rotate_doc_id": doc_id, "rotate_ranges_text": ranges_text, "rotate_degrees": degrees},
        )

    fp = UPLOAD_DIR / (doc.get("stored_name") or "")
    if not fp.exists():
        return _render_pdf_tools(
            request,
            status_code=404,
            result="âŒ A PDF fÃ¡jl nem talÃ¡lhatÃ³ a szerveren (hiÃ¡nyzÃ³ uploads fÃ¡jl).",
            result_kind="error",
            form={"rotate_doc_id": doc_id, "rotate_ranges_text": ranges_text, "rotate_degrees": degrees},
        )

    try:
        out = rotate_pages_pdf_bytes(fp.read_bytes(), ranges, int(degrees))
    except Exception:
        return _render_pdf_tools(
            request,
            status_code=400,
            result="âš ï¸ A forgatÃ¡s fokszÃ¡ma 90/180/270 legyen.",
            result_kind="warn",
            form={"rotate_doc_id": doc_id, "rotate_ranges_text": ranges_text, "rotate_degrees": degrees},
        )

    out_original = f"rotated_{degrees}_{doc.get('original_name','document.pdf')}"
    out_title = f"Rotated {degrees}Â° â€” {doc.get('title') or doc.get('original_name','PDF')}"

    new_id = _store_pdf_bytes_as_document(out, title=out_title, original_name=out_original, language=doc.get("language") or "auto")
    new_doc = get_document(new_id)

    return _render_pdf_tools(
        request,
        result=f"âœ… ForgatÃ¡s kÃ©sz ({degrees}Â°) â€” Ãºj dokumentum kÃ©szÃ¼lt.",
        result_kind="success",
        created_docs=[new_doc] if new_doc else [],
        form={"rotate_doc_id": doc_id, "rotate_ranges_text": ranges_text, "rotate_degrees": degrees},
    )




@app.post("/pdf-tools/extract-pages", response_class=HTMLResponse)
def pdf_tools_extract_pages(request: Request, doc_id: int = Form(...), ranges_text: str = Form("")):
    doc = get_document(doc_id)
    if not doc:
        return _render_pdf_tools(
            request,
            status_code=404,
            result="âŒ Nem talÃ¡lom ezt a dokumentumot.",
            result_kind="error",
            form={"extract_doc_id": doc_id, "extract_ranges_text": ranges_text},
        )

    try:
        ranges = parse_ranges(ranges_text)
    except Exception as e:
        return _render_pdf_tools(
            request,
            status_code=400,
            result=f"âš ï¸ HibÃ¡s tartomÃ¡ny. PÃ©lda: 1-2, 4. RÃ©szlet: {e}",
            result_kind="warn",
            form={"extract_doc_id": doc_id, "extract_ranges_text": ranges_text},
        )
    if not ranges:
        return _render_pdf_tools(
            request,
            status_code=400,
            result="âš ï¸ Adj meg oldaltartomÃ¡nyt (pÃ©lda: 1-2, 4).",
            result_kind="warn",
            form={"extract_doc_id": doc_id, "extract_ranges_text": ranges_text},
        )

    fp = UPLOAD_DIR / (doc.get("stored_name") or "")
    if not fp.exists():
        return _render_pdf_tools(
            request,
            status_code=404,
            result="âŒ A PDF fÃ¡jl nem talÃ¡lhatÃ³ a szerveren (hiÃ¡nyzÃ³ uploads fÃ¡jl).",
            result_kind="error",
            form={"extract_doc_id": doc_id, "extract_ranges_text": ranges_text},
        )

    try:
        out = extract_pages_pdf_bytes(fp.read_bytes(), ranges)
    except Exception as e:
        return _render_pdf_tools(
            request,
            status_code=400,
            result=f"âŒ Extract nem sikerÃ¼lt: {e}",
            result_kind="error",
            form={"extract_doc_id": doc_id, "extract_ranges_text": ranges_text},
        )

    original_out = f"extract_{doc.get('original_name','document.pdf')}"
    title_out = f"Extract â€” {doc.get('title') or doc.get('original_name','PDF')}"
    new_id = _store_pdf_bytes_as_document(out, title=title_out, original_name=original_out, language=doc.get("language") or "auto")
    new_doc = get_document(new_id)

    return _render_pdf_tools(
        request,
        result="âœ… KivÃ¡gott (extract) PDF elkÃ©szÃ¼lt.",
        result_kind="success",
        created_docs=[new_doc] if new_doc else [],
        form={"extract_doc_id": doc_id, "extract_ranges_text": ranges_text},
    )


@app.post("/pdf-tools/reorder", response_class=HTMLResponse)
def pdf_tools_reorder(request: Request, doc_id: int = Form(...), sequence_text: str = Form("")):
    doc = get_document(doc_id)
    if not doc:
        return _render_pdf_tools(
            request,
            status_code=404,
            result="âŒ Nem talÃ¡lom ezt a dokumentumot.",
            result_kind="error",
            form={"reorder_doc_id": doc_id, "reorder_sequence_text": sequence_text},
        )

    fp = UPLOAD_DIR / (doc.get("stored_name") or "")
    if not fp.exists():
        return _render_pdf_tools(
            request,
            status_code=404,
            result="âŒ A PDF fÃ¡jl nem talÃ¡lhatÃ³ a szerveren (hiÃ¡nyzÃ³ uploads fÃ¡jl).",
            result_kind="error",
            form={"reorder_doc_id": doc_id, "reorder_sequence_text": sequence_text},
        )

    data = fp.read_bytes()
    total = 0
    try:
        total = pdf_page_count(data)
    except Exception:
        total = 0

    try:
        seq = parse_page_sequence(sequence_text, total_pages=total or 10**9)
    except Exception as e:
        return _render_pdf_tools(
            request,
            status_code=400,
            result=f"âš ï¸ HibÃ¡s sorrend (pÃ©lda: 3,1,2,5-7). RÃ©szlet: {e}",
            result_kind="warn",
            form={"reorder_doc_id": doc_id, "reorder_sequence_text": sequence_text},
        )

    if not seq:
        return _render_pdf_tools(
            request,
            status_code=400,
            result="âš ï¸ Adj meg oldalsorrendet (pÃ©lda: 3,1,2,5-7).",
            result_kind="warn",
            form={"reorder_doc_id": doc_id, "reorder_sequence_text": sequence_text},
        )

    try:
        out = reorder_pages_pdf_bytes(data, seq)
    except Exception as e:
        return _render_pdf_tools(
            request,
            status_code=400,
            result=f"âŒ Reorder nem sikerÃ¼lt: {e}",
            result_kind="error",
            form={"reorder_doc_id": doc_id, "reorder_sequence_text": sequence_text},
        )

    original_out = f"reorder_{doc.get('original_name','document.pdf')}"
    title_out = f"Reorder â€” {doc.get('title') or doc.get('original_name','PDF')}"
    new_id = _store_pdf_bytes_as_document(out, title=title_out, original_name=original_out, language=doc.get("language") or "auto")
    new_doc = get_document(new_id)

    return _render_pdf_tools(
        request,
        result="âœ… Reorder PDF elkÃ©szÃ¼lt.",
        result_kind="success",
        created_docs=[new_doc] if new_doc else [],
        form={"reorder_doc_id": doc_id, "reorder_sequence_text": sequence_text},
    )


# ---------------- Study (C modul) ----------------
def _parse_int_or_none(x: str) -> int | None:
    try:
        s = (x or "").strip()
        return int(s) if s else None
    except Exception:
        return None


@app.get("/study", response_class=HTMLResponse)
def study_home(request: Request, msg: str = ""):
    total, due = get_study_counts()
    ctx = {
        "request": request,
        "docs": list_documents(),
        "total": total,
        "due": due,
        "msg": (msg or "").strip(),
    }
    ctx.update(_settings_context())
    return templates.TemplateResponse("study.html", ctx)


@app.post("/study/generate")
def study_generate(
    request: Request,
    doc: str = Form(""),
    include_notes: str = Form("1"),
    include_docs: str = Form(""),
):
    """Generate study cards from notes (default) and optionally from document text.

    NOTE: "document text" works only for text-based PDFs. Scanned PDFs will usually produce 0 cards.
    """
    doc_selected = _parse_int_or_none(doc)

    do_notes = (include_notes == "1")
    do_docs = (include_docs == "1")

    created_notes = 0
    skipped_notes = 0
    created_docs = 0
    skipped_docs = 0
    empty_docs = 0

    # --- Notes -> cards ---
    if do_notes:
        notes = list_notes(limit=500, document_id=doc_selected)
        for n in notes:
            note_doc_id = doc_selected if doc_selected is not None else (n.get("document_id") or None)
            pairs = extract_qa_pairs(n.get("body") or "")
            for q, a in pairs:
                cid, is_new = create_study_card(q, a, document_id=note_doc_id, note_id=n.get("id"))
                if cid is None:
                    continue
                if is_new:
                    created_notes += 1
                else:
                    skipped_notes += 1

    # --- Document search_text -> cards ---
    if do_docs:
        if doc_selected is not None:
            docs_to_use = [get_document(doc_selected)]
        else:
            # Safety limit: avoid generating a massive deck by accident
            docs_to_use = list_documents()[:8]

        for d in docs_to_use:
            if not d:
                continue
            body = (d.get("search_text") or "").strip()
            if not body:
                empty_docs += 1
                continue

            # Keep generation bounded
            pairs = extract_qa_pairs(body)[:300]
            for q, a in pairs:
                cid, is_new = create_study_card(q, a, document_id=int(d["id"]), note_id=None)
                if cid is None:
                    continue
                if is_new:
                    created_docs += 1
                else:
                    skipped_docs += 1

    parts = []
    if do_notes:
        parts.append(f"ðŸ“ notes: +{created_notes} | dup: {skipped_notes}")
    if do_docs:
        parts.append(f"ðŸ“„ docs: +{created_docs} | dup: {skipped_docs} | no-text: {empty_docs}")

    if not parts:
        msg = "âš ï¸ Semmi nem tÃ¶rtÃ©nt (legalÃ¡bb egy forrÃ¡st jelÃ¶lj be)."
    else:
        msg = "âœ… KÃ¡rtyÃ¡k generÃ¡lva â€” " + " | ".join(parts)

    return RedirectResponse(url=f"/study?msg={quote_plus(msg)}", status_code=303)


@app.get("/study/cards", response_class=HTMLResponse)
def study_cards_page(request: Request, q: str = "", doc: str = ""):
    doc_selected = _parse_int_or_none(doc)
    cards = list_study_cards(q=(q or "").strip(), document_id=doc_selected, limit=300)
    ctx = {
        "request": request,
        "q": (q or "").strip(),
        "docs": list_documents(),
        "doc_selected": doc_selected,
        "cards": cards,
    }
    ctx.update(_settings_context())
    return templates.TemplateResponse("study_cards.html", ctx)


@app.get("/study/cards/new", response_class=HTMLResponse)
def study_card_new(request: Request, doc: str = ""):
    doc_selected = _parse_int_or_none(doc)
    ctx = {
        "request": request,
        "docs": list_documents(),
        "doc_selected": doc_selected,
        "card": None,
    }
    ctx.update(_settings_context())
    return templates.TemplateResponse("study_card_edit.html", ctx)


@app.post("/study/cards/new")
def study_card_new_post(
    request: Request,
    document_id: str = Form(""),
    question: str = Form(""),
    answer: str = Form(""),
):
    doc_id = _parse_int_or_none(document_id)
    create_study_card(question, answer, document_id=doc_id, note_id=None)
    return RedirectResponse(url="/study/cards", status_code=303)


@app.get("/study/cards/{card_id}/edit", response_class=HTMLResponse)
def study_card_edit(request: Request, card_id: int):
    card = get_study_card(card_id)
    ctx = {
        "request": request,
        "docs": list_documents(),
        "doc_selected": card.get("document_id") if card else None,
        "card": card,
        "practice": 0,
    }
    ctx.update(_settings_context())
    return templates.TemplateResponse("study_card_edit.html", ctx)


@app.post("/study/cards/{card_id}/edit")
def study_card_edit_post(
    request: Request,
    card_id: int,
    document_id: str = Form(""),
    question: str = Form(""),
    answer: str = Form(""),
):
    doc_id = _parse_int_or_none(document_id)
    update_study_card(card_id, question, answer, doc_id)
    return RedirectResponse(url="/study/cards", status_code=303)


@app.post("/study/cards/{card_id}/delete")
def study_card_delete_post(request: Request, card_id: int):
    delete_study_card(card_id)
    return RedirectResponse(url="/study/cards", status_code=303)


@app.get("/study/session", response_class=HTMLResponse)
def study_session(request: Request, doc: str = "", show: int = 0, card_id: int = 0, practice: int = 0):
    doc_selected = _parse_int_or_none(doc)
    show_int = int(show or 0)
    practice_int = int(practice or 0)

    card = None

    # If we are revealing a card, keep the practice flag from the query param
    if show_int == 1 and int(card_id) > 0:
        card = get_study_card(int(card_id))
        ctx = {
            "request": request,
            "docs": list_documents(),
            "doc_selected": doc_selected,
            "show": show_int,
            "card": card,
            "practice": practice_int,
        }
        ctx.update(_settings_context())
        return templates.TemplateResponse("study_session.html", ctx)

    # Otherwise pick the next due
    card = get_next_due_card(doc_selected)

    # No due -> practice mode (random)
    if not card:
        card = get_random_card(doc_selected)
        practice_int = 1 if card else 0

    ctx = {
        "request": request,
        "docs": list_documents(),
        "doc_selected": doc_selected,
        "show": show_int,
        "card": card,
        "practice": practice_int,
    }
    ctx.update(_settings_context())
    return templates.TemplateResponse("study_session.html", ctx)


@app.post("/study/review")



@app.post("/study/review")
def study_review(
    request: Request,
    card_id: int = Form(...),
    correct: int = Form(0),
    doc: str = Form(""),
):
    doc_selected = _parse_int_or_none(doc)
    review_card(int(card_id), bool(int(correct) == 1), source="session")
    url = "/study/session"
    if doc_selected is not None:
        url += f"?doc={doc_selected}"
    return RedirectResponse(url=url, status_code=303)


@app.get("/study/quiz", response_class=HTMLResponse)
def study_quiz(request: Request, doc: str = ""):
    doc_selected = _parse_int_or_none(doc)
    practice = 0

    card = get_next_due_card(doc_selected)
    if not card:
        card = get_random_card(doc_selected)
        practice = 1 if card else 0

    options: list[str] = []
    if card:
        distractors = get_random_distractors(exclude_card_id=int(card["id"]), document_id=doc_selected, k=3)
        options = [card["answer"], *distractors]
        random.shuffle(options)

    ctx = {
        "request": request,
        "docs": list_documents(),
        "doc_selected": doc_selected,
        "card": card,
        "practice": practice,
        "answered": 0,
        "options": options,
        "picked": "",
        "is_correct": False,
    }
    ctx.update(_settings_context())
    return templates.TemplateResponse("study_quiz.html", ctx)


@app.post("/study/quiz/answer", response_class=HTMLResponse)
def study_quiz_answer(
    request: Request,
    card_id: int = Form(...),
    picked: str = Form(""),
    doc: str = Form(""),
    practice: int = Form(0),
):
    doc_selected = _parse_int_or_none(doc)
    card = get_study_card(int(card_id))
    if not card:
        return RedirectResponse(url="/study/quiz", status_code=303)

    picked2 = (picked or "").strip()
    correct = picked2 == (card.get("answer") or "").strip()
    review_card(int(card_id), bool(correct), source="quiz")

    ctx = {
        "request": request,
        "docs": list_documents(),
        "doc_selected": doc_selected,
        "card": card,
        "practice": int(practice or 0),
        "answered": 1,
        "options": [],
        "picked": picked2,
        "is_correct": bool(correct),
    }
    ctx.update(_settings_context())
    return templates.TemplateResponse("study_quiz.html", ctx)


@app.get("/study/stats", response_class=HTMLResponse)



@app.get("/study/stats", response_class=HTMLResponse)
def study_stats_page(request: Request, doc: str = ""):
    doc_selected = _parse_int_or_none(doc)
    s = study_stats(doc_selected)
    ctx = {
        "request": request,
        "docs": list_documents(),
        "doc_selected": doc_selected,
        "total": s["total"],
        "due": s["due"],
        "dist": s["dist"],
        "acc": s["acc"],
        "weak": s["weak"],
    }
    ctx.update(_settings_context())
    return templates.TemplateResponse("study_stats.html", ctx)


@app.get("/study/export/csv")
def study_export_csv(doc: str = ""):
    import csv
    doc_selected = _parse_int_or_none(doc)
    cards = list_study_cards(q="", document_id=doc_selected, limit=5000)

    import io as _io
    buf = _io.StringIO()
    w = csv.writer(buf)
    w.writerow(["card_id", "document", "box", "due_at", "question", "answer"])
    for c in cards:
        w.writerow([c.get('id'), c.get('document_title') or '', c.get('box') or '', c.get('due_at') or '', c.get('question') or '', c.get('answer') or ''])
    data = buf.getvalue().encode('utf-8')

    fname = 'study_cards.csv'
    if doc_selected is not None:
        fname = f'study_cards_doc_{doc_selected}.csv'
    return StreamingResponse(_io.BytesIO(data), media_type='text/csv; charset=utf-8', headers={
        'Content-Disposition': f'attachment; filename="{fname}"'
    })
